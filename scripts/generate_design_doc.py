from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "教育行业智能答疑助手项目设计文档.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size: int = 12) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size=18 if level == 1 else 15 if level == 2 else 13, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str = "", first_line: bool = True) -> None:
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, size=12)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=12)
        p.paragraph_format.line_spacing = 1.2


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = footer.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run = footer.add_run(" 页")
    set_run_font(run, size=10)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("《教育行业智能答疑助手》")
    set_run_font(run, size=24, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目设计文档")
    set_run_font(run, size=22, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    fields = [
        ("课程名称", "大数据系统基础课程设计"),
        ("项目名称", "教育行业智能答疑助手"),
        ("项目版本", "V1.0"),
        ("小组成员", "待补充"),
        ("指导教师", "待补充"),
        ("完成日期", str(date.today())),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, size=14)
    doc.add_page_break()


def add_manual_toc(doc: Document) -> None:
    add_heading(doc, "目    录", 1)
    entries = [
        "1. 需求分析",
        "1.1 项目概述",
        "1.2 项目需求说明",
        "1.3 性能要求与系统边界",
        "2. 概要设计",
        "2.1 功能综述",
        "2.2 系统功能模块",
        "2.3 各模块开发要求",
        "2.4 数据建模",
        "2.5 数据库表清单",
        "3. 详细设计",
        "3.1 系统功能设计",
        "3.2 项目编码",
        "4. 项目实施",
        "4.1 环境配置要求",
        "4.2 系统启动与接口说明",
        "4.3 测试与运行验证",
        "4.4 当前边界与后续优化",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        set_run_font(run, size=12)
    doc.add_page_break()


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    add_cover(doc)
    add_manual_toc(doc)

    add_heading(doc, "1. 需求分析", 1)
    add_heading(doc, "1.1 项目概述", 2)
    add_para(doc, "本项目旨在开发一套面向高校课程学习场景的教育行业智能答疑助手。系统以教师课件、教材片段、高考试卷、课程讲义等资料为知识来源，通过文档解析、OCR 识别、知识切分、向量检索和大语言模型问答技术，为学生提供基于课程资料的智能答疑服务。")
    add_para(doc, "与通用聊天机器人不同，本系统强调“基于资料回答”和“引用可追溯”。用户提出课程概念、知识点或题目解析类问题时，系统会先在课程知识库中检索相关片段，再结合检索结果生成回答，并返回引用来源、问答路由、置信度和会话编号。")
    add_para(doc, "当前版本定位为本地/单机 MVP，重点完成后端能力建设，包括课程管理、资料上传、知识库构建、RAG 问答接口、聊天记录保存和文档切分结果查看，为后续前端页面、教师管理端和更复杂的教学应用打基础。")

    add_heading(doc, "1.2 项目需求说明", 2)
    add_para(doc, "系统主要需求包括课程资料管理、知识库构建和智能答疑三个方面。")
    add_bullets(doc, [
        "课程管理：支持创建课程、查询课程列表和查看指定课程信息。",
        "资料上传：支持按课程上传 PDF、DOCX、TXT、Markdown 等常见教学资料。",
        "文档解析：支持 PDF 文本抽取、DOCX 段落和表格读取、普通文本读取；扫描型 PDF 可调用本地 OCR。",
        "知识切分：对试卷类资料进行章节、题号、选项、答案、解析等结构化识别，并生成适合检索的知识块。",
        "向量索引：将知识块写入 Chroma 向量库，默认使用本地 BAAI/bge-m3 向量模型。",
        "智能问答：根据问题类型执行课程问答、题目解析、概念解释或范围外提示。",
        "引用溯源：回答结果返回 citations 字段，包含文件名、页码、章节标题和 chunk_id。",
        "聊天记录：保存用户和助手消息，支持通过 session_id 查询历史消息。",
    ])

    add_heading(doc, "1.3 性能要求与系统边界", 2)
    add_bullets(doc, [
        "资料导入应能处理单门课程下多份文档，失败时记录错误状态并避免影响其他文档。",
        "检索问答应在课程过滤条件下返回相关知识块，默认 top_k 为 5。",
        "不配置聊天模型时，系统应降级为基于检索原文的保守摘要，避免编造课程内容。",
        "当前 MVP 不实现拍照搜题、自动判卷、复杂数学公式 LaTeX 结构化识别和完整前端页面。",
        "系统默认本地运行，OCR 和本地 embedding 不依赖云服务；LLM 可通过 OpenAI 或兼容接口配置。",
    ])

    add_heading(doc, "2. 概要设计", 1)
    add_heading(doc, "2.1 功能综述", 2)
    add_para(doc, "教育行业智能答疑助手采用后端 API 服务形式提供能力，整体流程为：课程创建 -> 资料上传 -> 文档解析 -> 试题/文本切分 -> 向量索引 -> 用户提问 -> 检索相关资料 -> 生成回答 -> 返回引用和置信度。")
    add_para(doc, "系统通过 SQLite 保存课程、文档、知识块和聊天记录等元数据，通过 Chroma 保存向量索引和知识块检索信息。FastAPI 提供 RESTful 接口，便于前端页面、脚本工具或接口文档页面调用。")

    add_heading(doc, "2.2 系统功能模块", 2)
    add_table(doc, ["模块名称", "主要功能", "对应实现"], [
        ["课程管理模块", "创建课程、查询课程、校验课程是否存在", "app/api/courses.py"],
        ["文档管理模块", "上传文件、保存文件路径、记录索引状态、重建索引", "app/api/documents.py"],
        ["文档解析与 OCR 模块", "解析 PDF、DOCX、TXT、Markdown；扫描 PDF 调用 RapidOCR", "app/rag/loaders.py、app/rag/ocr.py"],
        ["知识切分模块", "识别章节、题号、选项、答案、解析，生成 chunk", "app/rag/splitter.py"],
        ["向量检索模块", "本地或 OpenAI embedding、Chroma 持久化、按课程检索", "app/storage/vectorstore.py、app/rag/retriever.py"],
        ["智能问答模块", "问题路由、上下文构造、Prompt 调用、降级回答、引用生成", "app/rag/chains.py"],
        ["聊天记录模块", "创建会话、保存用户和助手消息、查询历史消息", "app/api/chat.py、app/models/db.py"],
    ])

    add_heading(doc, "2.3 各模块开发要求", 2)
    add_bullets(doc, [
        "课程和文档接口需要返回清晰的错误信息，例如课程不存在、文档不存在或解析失败。",
        "文档解析模块需要兼容不同来源的课程资料，对扫描型 PDF 自动走 OCR 链路。",
        "切分模块需要优先保留试题结构，在无法识别试题时使用通用文本切分作为兜底。",
        "向量模块需要支持本地 embedding 和 OpenAI 兼容 embedding 的切换，并在切换模型后重建索引。",
        "问答模块必须基于检索资料回答，资料不足时明确提示，不虚构引用。",
    ])

    add_heading(doc, "2.4 数据建模", 2)
    add_para(doc, "系统包含课程、文档、知识块、聊天会话和聊天消息五类核心实体。课程与文档是一对多关系，文档与知识块是一对多关系，聊天会话与聊天消息是一对多关系。知识块同时保存 SQLite 元数据和 Chroma 向量索引 id，以支持结构化管理和语义检索。")
    add_table(doc, ["实体", "说明", "关系"], [
        ["Course", "课程信息，包含课程名称、描述和创建时间", "一个课程可包含多个文档"],
        ["Document", "上传资料信息，包含文件名、路径、索引状态和错误信息", "一个文档可切分为多个知识块"],
        ["Chunk", "知识块信息，包含文本内容、页码、章节标题和元数据", "每个知识块属于一个文档和课程"],
        ["ChatSession", "聊天会话信息，保存会话 id 和课程范围", "一个会话包含多条消息"],
        ["ChatMessage", "聊天消息信息，保存角色、内容、引用、检索片段和路由", "每条消息属于一个会话"],
    ])

    add_heading(doc, "2.5 数据库表清单", 2)
    add_table(doc, ["表名", "字段名", "数据类型", "描述", "约束条件"], [
        ["courses", "id", "Integer", "课程 ID", "主键，自增"],
        ["courses", "name", "String(200)", "课程名称", "非空"],
        ["courses", "description", "Text", "课程描述", "默认空字符串"],
        ["courses", "created_at", "DateTime", "创建时间", "默认当前时间"],
        ["documents", "id", "Integer", "文档 ID", "主键，自增"],
        ["documents", "course_id", "Integer", "所属课程 ID", "外键，索引"],
        ["documents", "file_name", "String(300)", "原始文件名", "非空"],
        ["documents", "source_path", "Text", "文件保存路径", "非空"],
        ["documents", "status", "String(50)", "索引状态", "pending/indexed/failed"],
        ["documents", "error_message", "Text", "失败错误信息", "默认空字符串"],
        ["chunks", "chunk_id", "String(120)", "知识块业务 ID", "唯一，索引"],
        ["chunks", "content", "Text", "知识块正文", "非空"],
        ["chunks", "page", "Integer", "来源页码", "可为空"],
        ["chunks", "section_title", "String(300)", "章节标题", "默认空字符串"],
        ["chunks", "metadata_json", "JSON", "知识块元数据", "默认空对象"],
        ["chat_sessions", "id", "String(80)", "会话 ID", "主键"],
        ["chat_sessions", "course_id", "Integer", "会话绑定课程 ID", "可为空，索引"],
        ["chat_messages", "role", "String(30)", "消息角色", "user/assistant"],
        ["chat_messages", "citations", "JSON", "回答引用信息", "默认空列表"],
        ["chat_messages", "retrieved_chunks", "JSON", "检索到的知识块", "默认空列表"],
    ])

    add_heading(doc, "3. 详细设计", 1)
    add_heading(doc, "3.1 系统功能设计", 2)
    add_heading(doc, "3.1.1 课程管理设计", 3)
    add_para(doc, "课程管理模块提供课程创建、课程列表查询和单个课程查询能力。创建课程时系统保存课程名称、描述和创建时间；上传文档和发起问答时会校验课程是否存在，避免无效数据进入知识库。")
    add_heading(doc, "3.1.2 文档导入设计", 3)
    add_para(doc, "用户通过 /api/courses/{course_id}/documents 上传资料。系统先将文件保存到本地上传目录，再在 documents 表中创建 pending 状态记录，随后调用 ingest_document 执行解析、切分和索引。若处理成功，状态更新为 indexed；若出现异常，状态更新为 failed 并写入 error_message。")
    add_heading(doc, "3.1.3 文档解析设计", 3)
    add_para(doc, "文档解析模块根据文件后缀选择不同处理方式。PDF 先使用 pypdf 抽取文本，如果判断为扫描型 PDF，则调用 pdftoppm 将页面渲染为图片，再使用 RapidOCR 识别文字。DOCX 文件通过 python-docx 读取段落和表格，同时从 OOXML 关系中提取图片引用。TXT、Markdown 等文本文件直接按 UTF-8 读取。")
    add_heading(doc, "3.1.4 试题切分设计", 3)
    add_para(doc, "切分模块针对高考解析卷等资料设计，先清理页码、网站水印等噪声，再识别章节标题和题号，将题干、选项、答案、解析、点评、图片和公式标记拆分成结构化对象。对于无法识别为试题的普通资料，系统使用 RecursiveCharacterTextSplitter 按 chunk_size 和 chunk_overlap 做通用分块。")
    add_heading(doc, "3.1.5 向量知识库设计", 3)
    add_para(doc, "向量知识库使用 Chroma PersistentClient 持久化保存。默认 embedding_provider 为 local，使用 BAAI/bge-m3 模型在 CPU 上生成向量；也可以配置 EMBEDDING_PROVIDER=openai 使用 OpenAI 兼容 embedding。每个知识块写入 Chroma 时携带 course_id、document_id、file_name、page、section_title、question_no、question_type、quality_flags 等元数据。")
    add_heading(doc, "3.1.6 智能问答设计", 3)
    add_para(doc, "问答模块先根据关键词判断问题路由，包括 knowledge_qa、problem_solving、concept_explain 和 out_of_scope。非范围外问题会调用 retrieve_chunks 检索相关资料，再根据路由选择课程问答 Prompt 或题目解析 Prompt。若未配置聊天模型或没有检索结果，系统返回保守摘要或资料不足提示。")
    add_heading(doc, "3.1.7 引用与聊天记录设计", 3)
    add_para(doc, "回答结果中的 citations 来源于检索知识块元数据，包含文件名、页码、章节标题和 chunk_id。系统会把用户问题和助手回答分别保存到 chat_messages 表中，并通过 session_id 关联为同一次对话，便于后续查看历史记录。")

    add_heading(doc, "3.2 项目编码", 2)
    add_para(doc, "以下选取核心代码片段说明系统关键实现。")
    add_para(doc, "（1）FastAPI 应用入口：")
    add_code(doc, """
app = FastAPI(title=settings.app_name)

@app.on_event("startup")
def on_startup() -> None:
    init_db()

@app.get("/api/health")
def health_check():
    return {"status": "ok", "app": settings.app_name}

app.include_router(courses_router)
app.include_router(documents_router)
app.include_router(chat_router)
""")
    add_para(doc, "（2）文档上传与索引：")
    add_code(doc, """
path = await save_upload_file(course_id, file)
document = Document(
    course_id=course_id,
    file_name=file.filename or path.name,
    source_path=str(path),
    status="pending",
)
db.add(document)
db.commit()
db.refresh(document)
ingest_document(db, document)
""")
    add_para(doc, "（3）文档入库主流程：")
    add_code(doc, """
pages = load_text_pages(path)
chunks = split_documents(pages)
collection = get_chroma_collection()

for index, item in enumerate(chunks, start=1):
    chunk_id = f"doc-{document.id}-{index}-{uuid4().hex[:8]}"
    metadata = {
        "course_id": document.course_id,
        "document_id": document.id,
        "file_name": document.file_name,
        "page": item.get("page"),
        "section_title": item.get("section_title", ""),
        "chunk_id": chunk_id,
    }
collection.add(ids=ids, documents=texts, metadatas=metadatas)
""")
    add_para(doc, "（4）向量检索：")
    add_code(doc, """
where = {"course_id": course_id} if course_id is not None else None
results = collection.query(
    query_texts=[question],
    n_results=top_k or settings.retrieval_top_k,
    where=where,
)
""")
    add_para(doc, "（5）问答路由和降级回答：")
    add_code(doc, """
route = route_question(question)
chunks = retrieve_chunks(question, course_id=course_id)
citations = citations_from_chunks(chunks)
confidence = estimate_confidence(chunks)
llm = get_chat_llm()

if llm is None or not chunks:
    answer = fallback_answer(question, route, chunks)
else:
    context = build_context(chunks)
    answer = llm.invoke(prompt_template.format(
        question=question,
        context=context,
    )).content
""")

    add_heading(doc, "4. 项目实施", 1)
    add_heading(doc, "4.1 环境配置要求", 2)
    add_table(doc, ["类别", "技术/依赖", "作用"], [
        ["Web 框架", "FastAPI、Uvicorn", "提供 RESTful API 和接口文档"],
        ["数据存储", "SQLite、SQLAlchemy", "保存课程、文档、知识块和聊天记录元数据"],
        ["向量数据库", "ChromaDB", "保存文档向量索引并执行相似度检索"],
        ["RAG 框架", "LangChain、langchain-openai", "Prompt 构造和大模型调用"],
        ["Embedding", "sentence-transformers、BAAI/bge-m3", "本地文本向量化"],
        ["文档解析", "pypdf、python-docx", "读取 PDF 和 DOCX 文档内容"],
        ["OCR", "RapidOCR、pdftoppm/Poppler", "扫描 PDF 页面渲染与文字识别"],
        ["配置管理", "pydantic-settings、.env", "管理模型、数据库、目录和检索参数"],
    ])
    add_para(doc, "安装依赖命令如下：")
    add_code(doc, """
conda activate vec_shaorixuan
pip install -r requirements.txt
Copy-Item .env.example .env
""")

    add_heading(doc, "4.2 系统启动与接口说明", 2)
    add_para(doc, "系统启动命令如下：")
    add_code(doc, "uvicorn app.main:app --host 127.0.0.1 --port 8000 --reload")
    add_para(doc, "启动后可以访问 http://127.0.0.1:8000/docs 查看 FastAPI 自动生成的接口文档。主要接口如下：")
    add_table(doc, ["接口", "方法", "说明"], [
        ["/api/health", "GET", "健康检查"],
        ["/api/courses", "POST", "创建课程"],
        ["/api/courses", "GET", "查看课程列表"],
        ["/api/courses/{course_id}", "GET", "查看课程详情"],
        ["/api/courses/{course_id}/documents", "POST", "上传并索引课程资料"],
        ["/api/courses/{course_id}/reindex", "POST", "重建课程文档索引"],
        ["/api/chat", "POST", "课程智能问答"],
        ["/api/chat/{session_id}/messages", "GET", "查看聊天记录"],
        ["/api/documents/{document_id}/chunks", "GET", "查看文档切分结果"],
    ])

    add_heading(doc, "4.3 测试与运行验证", 2)
    add_para(doc, "测试时可选择 Math 课程作为示例，先创建课程，再上传 data/uploads/Math 目录下的高考试卷解析卷，确认系统完成解析和索引。随后通过 /api/documents/{document_id}/chunks 查看切分结果，通过 /api/chat 提交概念解释或题目解析类问题，检查返回 answer、citations、route、confidence 和 session_id 字段。")
    add_bullets(doc, [
        "课程创建测试：调用 POST /api/courses，确认返回课程 ID。",
        "文档上传测试：上传 DOCX 或 PDF，确认 documents.status 最终为 indexed。",
        "切分结果测试：查看 chunks，确认存在 section_title、question_no、question_type 等元数据。",
        "问答测试：提出课程相关问题，确认系统返回引用来源和置信度。",
        "边界测试：提出天气、股票等范围外问题，确认 route 为 out_of_scope。",
        "降级测试：不配置 OPENAI_API_KEY 时，确认系统返回基于资料的保守摘要。",
    ])

    add_heading(doc, "4.4 当前边界与后续优化", 2)
    add_para(doc, "当前版本已经完成教育智能答疑助手的核心后端 MVP，但仍存在可优化空间。")
    add_bullets(doc, [
        "前端页面尚未实现，目前主要通过 FastAPI Swagger 页面或脚本调用接口。",
        "数学公式和复杂版式主要以文本或图片标记保留，暂未完整转换为 LaTeX。",
        "部分历史实验模块中存在编码异常文本，正式交付前建议清理无关模块或修正文案。",
        "后续可增加教师端资料管理、学生端问答页面、问答评价、知识库统计和权限控制。",
        "后续可补充接口截图、问答截图、项目目录截图和答辩用流程图，使文档展示效果更完整。",
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录：截图素材建议", 1)
    add_para(doc, "为便于课程答辩，建议在最终版文档中补充以下截图：FastAPI /docs 页面、创建课程接口、上传文档接口、问答接口返回结果、项目目录结构、data/uploads 示例资料目录。")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
